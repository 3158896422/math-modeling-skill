import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from lxml import etree


SCRIPTS = Path(__file__).resolve().parents[1] / "tools" / "docx" / "scripts"
sys.path.insert(0, str(SCRIPTS))

from equations import OMML_NS, latex2omml, latex_to_docx, replace_placeholder


class EquationConversionTests(unittest.TestCase):
    def test_physics_greek_letters_relations_and_inverse_trig_are_supported(self):
        xml = etree.fromstring(
            latex2omml(r"\nu \approx \mu,\quad \theta=\arcsin(x)+\arccos(y)+\arctan(z)")
        )
        text = "".join(xml.itertext())
        for expected in ("ν", "≈", "μ", "arcsin", "arccos", "arctan"):
            self.assertIn(expected, text)

    def test_common_physics_constants_delimiters_and_number_sets_are_supported(self):
        xml = etree.fromstring(
            latex2omml(r"\hbar\omega,\quad \ell\in\mathbb{R},\quad \langle x\rangle=30^\circ")
        )
        text = "".join(xml.itertext())
        for expected in ("ℏ", "ω", "ℓ", "∈", "R", "⟨", "⟩", "°"):
            self.assertIn(expected, text)

    def test_common_commands_are_not_silently_corrupted(self):
        xml = etree.fromstring(latex2omml(r"x_1,\ldots,x_n"))
        text = "".join(xml.itertext())
        self.assertIn("…", text)
        self.assertNotIn("ldots", text)

    def test_nth_root_contains_visible_degree(self):
        xml = etree.fromstring(latex2omml(r"\sqrt[3]{x}"))
        degree = xml.find(f".//{{{OMML_NS}}}deg")
        self.assertIsNotNone(degree)
        self.assertIn("3", "".join(degree.itertext()))

    def test_matrix_generates_matrix_rows(self):
        xml = etree.fromstring(latex2omml(r"\begin{bmatrix}a&b\\c&d\end{bmatrix}"))
        rows = xml.findall(f".//{{{OMML_NS}}}mr")
        self.assertEqual(len(rows), 2)

    def test_unknown_command_fails_explicitly(self):
        with self.assertRaisesRegex(ValueError, "不支持"):
            latex2omml(r"\unknowncommand{x}")

    def test_unbalanced_group_fails_explicitly(self):
        with self.assertRaisesRegex(ValueError, "括号"):
            latex2omml(r"x_{i")
        with self.assertRaisesRegex(ValueError, "括号"):
            latex2omml("x}")

    def test_common_modeling_operators_and_styles_are_supported(self):
        xml = etree.fromstring(
            latex2omml(r"\min_x \operatorname{RMSE}(x)+\log(x)+\mathbf{w}^T x")
        )
        text = "".join(xml.itertext())
        for expected in ("min", "RMSE", "log", "w"):
            self.assertIn(expected, text)

    def test_cases_environment_generates_two_rows(self):
        xml = etree.fromstring(
            latex2omml(r"\begin{cases}x,&x>0\\0,&x\le 0\end{cases}")
        )
        rows = xml.findall(f".//{{{OMML_NS}}}mr")
        self.assertEqual(len(rows), 2)

    def test_replaces_every_matching_placeholder(self):
        doc = Document()
        doc.add_paragraph("第一处 EQ")
        doc.add_paragraph("第二处 EQ")

        replaced = replace_placeholder(doc, "EQ", "x")

        self.assertEqual(replaced, 2)
        self.assertFalse(any("EQ" in p.text for p in doc.paragraphs))

    def test_replaces_placeholder_inside_table_cell(self):
        doc = Document()
        cell = doc.add_table(rows=1, cols=1).cell(0, 0)
        cell.text = "指标 EQ"

        replaced = replace_placeholder(doc, "EQ", "x")

        self.assertEqual(replaced, 1)
        self.assertNotIn("EQ", cell.text)

    def test_converts_complete_latex_to_docx_with_pandoc_atomically(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "main.tex"
            source.write_text(
                r"\documentclass{article}\begin{document}$x^2$\end{document}",
                encoding="utf-8",
            )
            template = root / "template.docx"
            Document().save(template)
            output = root / "paper.docx"
            observed = []

            def fake_run(command, **kwargs):
                observed.append((command, kwargs))
                generated = Path(command[command.index("--output") + 1])
                document = Document()
                document.add_paragraph("converted")
                document.save(generated)
                return subprocess.CompletedProcess(
                    command, 0, stdout="", stderr="conversion warning"
                )

            with patch("equations.shutil.which", return_value="pandoc"), patch(
                "equations.subprocess.run", side_effect=fake_run
            ):
                result = latex_to_docx(source, output, template)

            command, options = observed[0]
            self.assertEqual(command[command.index("--from") + 1], "latex")
            self.assertIn("--citeproc", command)
            self.assertIn(f"--resource-path={source.parent}", command)
            self.assertIn("--reference-doc", command)
            self.assertEqual(options["cwd"], source.parent)
            self.assertTrue(output.is_file())
            self.assertEqual(result["warnings"], ["conversion warning"])
            self.assertEqual(
                {path.name for path in root.iterdir()},
                {"main.tex", "template.docx", "paper.docx"},
            )

            with self.assertRaises(FileExistsError):
                latex_to_docx(source, output, template)


if __name__ == "__main__":
    unittest.main()
